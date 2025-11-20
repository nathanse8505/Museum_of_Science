# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

HEB_IMAGES_DIR = "תמונות"
SPEC_FILENAME = "מפרט טכני.docx"
H1_TEXT = "ציוד"

def create_docx_if_missing(doc_path: Path, folder_name: str):
    """Crée le fichier Word uniquement s'il n'existe pas déjà."""
    if doc_path.exists():
        # fichier déjà présent → NE PAS ÉCRASER
        return

    doc = Document()
    doc.core_properties.title = folder_name

    # Titre (nom du dossier)
    title_para = doc.add_paragraph(folder_name, style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")  # espace

    # Heading 1 = ציוד
    h1_para = doc.add_paragraph(H1_TEXT)
    h1_para.style = doc.styles["Heading 1"]
    h1_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(str(doc_path))


def process_parent(parent: Path):
    if not parent.exists() or not parent.is_dir():
        messagebox.showerror("Erreur", f"Le dossier n'existe pas : {parent}")
        return

    count = 0

    for child in sorted(parent.iterdir()):
        if not child.is_dir() or child.name.startswith("."):
            continue

        # Créer le dossier תמונות S'IL N'EXISTE PAS
        images_dir = child / HEB_IMAGES_DIR
        if not images_dir.exists():
            images_dir.mkdir(parents=True, exist_ok=True)

        # Créer le fichier מפרט טכני S'IL N'EXISTE PAS
        doc_path = child / SPEC_FILENAME
        create_docx_if_missing(doc_path, folder_name=child.name)

        count += 1

    messagebox.showinfo("Terminé", f"{count} sous-dossiers traités avec succès.")


def choisir_dossier():
    chemin = filedialog.askdirectory(title="Choisir le dossier parent")
    if chemin:
        entree_chemin.delete(0, tk.END)
        entree_chemin.insert(0, chemin)


def lancer_creation():
    chemin = entree_chemin.get().strip()
    if not chemin:
        messagebox.showwarning("Attention", "Merci de saisir ou choisir un dossier parent.")
        return
    process_parent(Path(chemin))


# --- Interface graphique ---
root = tk.Tk()
root.title("Création des dossiers 'תמונות' et fichiers 'מפרט טכני'")
root.geometry("600x180")

tk.Label(root, text="Chemin du dossier parent :", font=("Segoe UI", 11)).pack(pady=10)
frame = tk.Frame(root)
frame.pack()

entree_chemin = tk.Entry(frame, width=60, font=("Segoe UI", 10))
entree_chemin.pack(side=tk.LEFT, padx=5)
tk.Button(frame, text="Parcourir...", command=choisir_dossier).pack(side=tk.LEFT)

tk.Button(root, text="Créer dans tous les sous-dossiers", font=("Segoe UI", 11, "bold"), bg="#4CAF50", fg="white",
          command=lancer_creation).pack(pady=15)

root.mainloop()
